"""
Word 文档高亮文本提取工具
从 Word 文档中提取黄色高亮文本，并按规章条文整合成填空题
"""

from docx import Document
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import re
import sys


def get_highlight_color(run):
    """
    获取文本的高亮颜色
    返回颜色索引，如果是黄色返回 True
    """
    try:
        # 获取高亮属性
        highlight_elm = run._element.find(qn('w:highlight'))
        if highlight_elm is not None:
            highlight_val = highlight_elm.get(qn('w:val'))
            # 黄色的值为 'yellow'
            return highlight_val == 'yellow'
    except:
        pass
    return False


def extract_highlights_from_paragraph(paragraph):
    """
    从段落中提取高亮文本和原始文本（用下划线替换高亮部分）
    返回 (原始文本, 高亮文本列表)
    """
    highlighted_texts = []
    modified_text = ""
    
    for run in paragraph.runs:
        if get_highlight_color(run):
            highlighted_texts.append(run.text)
            # 用下划线替换高亮文本
            modified_text += "_" * len(run.text)
        else:
            modified_text += run.text
    
    return modified_text, highlighted_texts


def extract_articles(docx_path):
    """
    从 Word 文档中提取所有规章条文和高亮内容
    返回字典，key 为条号，value 为 (原始文本, 高亮答案列表)
    """
    doc = Document(docx_path)
    articles = {}
    current_article_num = None
    current_content = ""
    current_highlights = []
    
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        
        if not text:
            continue
        
        # 检测新的条文开始 (如 "第 247 条" 或 "第247条")
        article_match = re.match(r'第\s*(\d+)\s*条', text)
        
        if article_match:
            # 保存前一个条文
            if current_article_num is not None:
                articles[current_article_num] = (current_content, current_highlights)
            
            # 开始新条文
            current_article_num = article_match.group(1)
            current_content = text
            current_highlights = []
            
            # 从该段落提取高亮
            _, highlights = extract_highlights_from_paragraph(paragraph)
            current_highlights.extend(highlights)
        else:
            # 继续添加到当前条文
            if current_article_num is not None:
                current_content += "\n" + text
                _, highlights = extract_highlights_from_paragraph(paragraph)
                current_highlights.extend(highlights)
    
    # 保存最后一个条文
    if current_article_num is not None:
        articles[current_article_num] = (current_content, current_highlights)
    
    return articles


def generate_output_document(articles, output_path):
    """
    生成输出 Word 文档，格式为填空题
    """
    doc = Document()
    
    # 添加标题
    title = doc.add_heading('填空题提取结果', level=1)
    
    for article_num in sorted(articles.keys(), key=lambda x: int(x)):
        content, highlights = articles[article_num]
        
        if not highlights:
            continue
        
        # 添加条文号和内容
        doc.add_heading(f'第 {article_num} 条', level=2)
        
        # 添加题目
        doc.add_paragraph(f'题目：{content}', style='List Bullet')
        
        # 添加答案
        answers_text = '、'.join(highlights)
        doc.add_paragraph(f'答案：{answers_text}', style='List Bullet')
        
        # 添加空行
        doc.add_paragraph()
    
    doc.save(output_path)
    print(f"✓ 输出文件已生成：{output_path}")


def main():
    if len(sys.argv) < 2:
        print("使用方法：python extract_highlights.py <input.docx> [output.docx]")
        print("示例：python extract_highlights.py input.docx output.docx")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else "output.docx"
    
    try:
        print(f"正在处理文件：{input_file}")
        articles = extract_articles(input_file)
        
        if not articles:
            print("✗ 未找到任何规章条文或高亮内容")
            sys.exit(1)
        
        print(f"✓ 找到 {len(articles)} 条规章")
        
        generate_output_document(articles, output_file)
        print(f"✓ 处理完成！")
        
    except FileNotFoundError:
        print(f"✗ 错误：找不到文件 {input_file}")
        sys.exit(1)
    except Exception as e:
        print(f"✗ 错误：{str(e)}")
        sys.exit(1)


if __name__ == "__main__":
    main()
